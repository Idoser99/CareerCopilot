"""ATS-friendly DOCX rendering helpers for the CV writer tool."""

from pathlib import Path
from typing import Any, Mapping, Sequence, cast

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.styles.style import ParagraphStyle


STANDARD_SECTION_TITLES = {
    "professional_summary": "Professional Summary",
    "education": "Education",
    "technical_skills": "Technical Skills",
    "experience": "Experience",
    "projects": "Projects",
    "certifications": "Certifications",
    "awards": "Awards",
    "leadership": "Leadership",
    "languages": "Languages",
}


def _set_style_font(
    style: ParagraphStyle,
    font_name: str,
    size_pt: float,
    **attributes: Any,
) -> None:
    """Set a style font explicitly for consistent Word-compatible rendering."""
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    color = attributes.pop("color", None)
    if color is not None:
        style.font.color.rgb = color
    for attribute, value in attributes.items():
        setattr(style.font, attribute, value)

    style_element = style._element
    if style_element is None:
        raise ValueError("Paragraph style is missing its XML element")

    run_properties = style_element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    for font_attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        run_fonts.set(qn(font_attribute), font_name)


def _paragraph_style(
    document: DocumentType,
    name: str,
    base_style: str = "Normal",
) -> ParagraphStyle:
    styles = document.styles
    if name in styles:
        return cast(ParagraphStyle, styles[name])
    style = cast(
        ParagraphStyle,
        styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH),
    )
    style.base_style = cast(ParagraphStyle, styles[base_style])
    return style


def _configure_page(document: DocumentType, preferences: Mapping[str, Any]) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT

    if preferences["page_size"] == "LETTER":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    else:
        section.page_width = Mm(210)
        section.page_height = Mm(297)

    margin = Inches(preferences["margin_inches"])
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def _configure_styles(document: DocumentType, preferences: Mapping[str, Any]) -> None:
    styles = document.styles
    font_name = preferences["font_name"]
    body_size = preferences["font_size_pt"]

    normal = cast(ParagraphStyle, styles["Normal"])
    _set_style_font(normal, font_name, body_size, color=RGBColor(0, 0, 0))
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.widow_control = True

    heading = cast(ParagraphStyle, styles["Heading 1"])
    heading.base_style = normal
    _set_style_font(
        heading,
        font_name,
        body_size + 1,
        bold=True,
        color=RGBColor(0, 0, 0),
    )
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True

    list_bullet = cast(ParagraphStyle, styles["List Bullet"])
    list_bullet.base_style = normal
    _set_style_font(list_bullet, font_name, body_size, color=RGBColor(0, 0, 0))
    list_bullet.paragraph_format.left_indent = Inches(0.22)
    list_bullet.paragraph_format.first_line_indent = Inches(-0.14)
    list_bullet.paragraph_format.space_after = Pt(1.5)
    list_bullet.paragraph_format.line_spacing = 1.0

    name_style = _paragraph_style(document, "CV Name")
    _set_style_font(
        name_style,
        font_name,
        19,
        bold=True,
        color=RGBColor(0, 0, 0),
    )
    name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_style.paragraph_format.space_after = Pt(2)
    name_style.paragraph_format.keep_with_next = True

    contact_style = _paragraph_style(document, "CV Contact")
    _set_style_font(contact_style, font_name, 9.5, color=RGBColor(0, 0, 0))
    contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_style.paragraph_format.space_after = Pt(1)
    contact_style.paragraph_format.line_spacing = 1.0
    contact_style.paragraph_format.keep_with_next = True

    entry_style = _paragraph_style(document, "CV Entry Title")
    _set_style_font(
        entry_style,
        font_name,
        body_size,
        bold=True,
        color=RGBColor(0, 0, 0),
    )
    entry_style.paragraph_format.space_before = Pt(3)
    entry_style.paragraph_format.space_after = Pt(0)
    entry_style.paragraph_format.keep_with_next = True

    details_style = _paragraph_style(document, "CV Entry Details")
    _set_style_font(
        details_style,
        font_name,
        max(9.5, body_size - 0.5),
        italic=True,
        color=RGBColor(0, 0, 0),
    )
    details_style.paragraph_format.space_after = Pt(1)
    details_style.paragraph_format.keep_with_next = True


def _create_document(preferences: Mapping[str, Any]) -> DocumentType:
    document = Document()
    _configure_page(document, preferences)
    _configure_styles(document, preferences)
    return document


def _add_contact_block(document: DocumentType, contact: Mapping[str, Any]) -> None:
    name_paragraph = document.add_paragraph(style="CV Name")
    name_paragraph.add_run(contact["full_name"])

    main_details = [
        contact.get("location"),
        contact.get("email"),
        contact.get("phone"),
    ]
    links = [
        contact.get("linkedin"),
        contact.get("github"),
        contact.get("portfolio"),
    ]
    detail_lines = [
        " | ".join(value for value in values if value)
        for values in (main_details, links)
        if any(values)
    ]

    for line in detail_lines:
        document.add_paragraph(line, style="CV Contact")

    if detail_lines:
        document.paragraphs[-1].paragraph_format.space_after = Pt(4)
    else:
        name_paragraph.paragraph_format.space_after = Pt(4)


def _add_section_heading(document: DocumentType, title: str) -> None:
    document.add_paragraph(title, style="Heading 1")


def _add_professional_summary(document: DocumentType, summary: str) -> None:
    paragraph = document.add_paragraph(summary, style="Normal")
    paragraph.paragraph_format.space_after = Pt(3)


def _add_entry(document: DocumentType, entry: Mapping[str, Any]) -> None:
    title_parts = [entry["primary_text"]]
    if entry.get("secondary_text"):
        title_parts.append(entry["secondary_text"])
    title = " - ".join(title_parts)
    title_paragraph = document.add_paragraph(title, style="CV Entry Title")

    details = [
        entry.get("location"),
        entry.get("date"),
        entry.get("link"),
    ]
    details_text = " | ".join(value for value in details if value)
    bullets = entry.get("bullets") or []

    if details_text:
        document.add_paragraph(details_text, style="CV Entry Details")

    for bullet in bullets:
        document.add_paragraph(bullet, style="List Bullet")

    last_entry_paragraph = (
        document.paragraphs[-1] if details_text or bullets else title_paragraph
    )
    last_entry_paragraph.paragraph_format.space_after = Pt(3)
    last_entry_paragraph.paragraph_format.keep_with_next = False


def _add_entries(document: DocumentType, entries: Sequence[Mapping[str, Any]]) -> None:
    for entry in entries:
        _add_entry(document, entry)


def _add_skill_groups(
    document: DocumentType, groups: Sequence[Mapping[str, Any]]
) -> None:
    for group in groups:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.add_run(f"{group['label']}: ").bold = True
        paragraph.add_run(", ".join(group["items"]))


def _add_languages(document: DocumentType, languages: Sequence[str]) -> None:
    for language in languages:
        document.add_paragraph(language, style="List Bullet")


def _add_custom_block(document: DocumentType, block: Mapping[str, Any]) -> None:
    block_type = block["block_type"]
    if block_type == "paragraph":
        document.add_paragraph(block["text"], style="Normal")
    elif block_type == "labeled_text":
        paragraph = document.add_paragraph(style="Normal")
        paragraph.add_run(f"{block['label']}: ").bold = True
        paragraph.add_run(block["text"])
    elif block_type == "entry":
        _add_entry(document, block["entry"])
    elif block_type == "bullet_list":
        for bullet in block["bullets"]:
            document.add_paragraph(bullet, style="List Bullet")
    else:
        raise ValueError(f"Unsupported custom block type: {block_type}")


def _add_custom_section(
    document: DocumentType, custom_section: Mapping[str, Any]
) -> None:
    for block in custom_section["blocks"]:
        _add_custom_block(document, block)


def _custom_sections_by_id(payload: Mapping[str, Any]) -> dict[str, Mapping[str, Any]]:
    return {
        section["section_id"]: section
        for section in payload.get("custom_sections", [])
    }


def _section_title(
    section_id: str, custom_sections: Mapping[str, Mapping[str, Any]]
) -> str:
    if section_id in STANDARD_SECTION_TITLES:
        return STANDARD_SECTION_TITLES[section_id]
    return custom_sections[section_id]["title"]


def _render_section(
    document: DocumentType,
    section_id: str,
    payload: Mapping[str, Any],
    custom_sections: Mapping[str, Mapping[str, Any]],
) -> None:
    _add_section_heading(document, _section_title(section_id, custom_sections))

    if section_id == "professional_summary":
        _add_professional_summary(document, payload[section_id])
    elif section_id == "technical_skills":
        _add_skill_groups(document, payload[section_id])
    elif section_id == "languages":
        _add_languages(document, payload[section_id])
    elif section_id in STANDARD_SECTION_TITLES:
        _add_entries(document, payload[section_id])
    else:
        _add_custom_section(document, custom_sections[section_id])


def render_cv_docx(
    payload: Mapping[str, Any],
    ordered_section_ids: Sequence[str],
    output_path: Path,
) -> None:
    """Construct and save an ATS-friendly CV without changing supplied content."""
    document = _create_document(payload["document_preferences"])
    _add_contact_block(document, payload["contact"])

    custom_sections = _custom_sections_by_id(payload)
    for section_id in ordered_section_ids:
        _render_section(document, section_id, payload, custom_sections)

    document.core_properties.title = f"{payload['contact']['full_name']} CV"
    document.core_properties.subject = "Curriculum Vitae"
    document.save(str(output_path))


def verify_generated_docx(
    output_path: Path,
    payload: Mapping[str, Any],
    ordered_section_ids: Sequence[str],
) -> None:
    """Reopen a generated file and verify its essential document structure."""
    document = Document(str(output_path))
    paragraph_texts = [paragraph.text for paragraph in document.paragraphs]

    if len(document.sections) != 1:
        raise ValueError("Generated CV must contain exactly one document section")
    if document.tables:
        raise ValueError("Generated CV must not use layout tables")
    if payload["contact"]["full_name"] not in paragraph_texts:
        raise ValueError("Generated CV is missing the candidate name")
    if payload["professional_summary"] not in paragraph_texts:
        raise ValueError("Generated CV is missing the professional summary")

    custom_sections = _custom_sections_by_id(payload)
    expected_headings = [
        _section_title(section_id, custom_sections)
        for section_id in ordered_section_ids
    ]
    actual_headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style and paragraph.style.name == "Heading 1"
    ]
    if actual_headings != expected_headings:
        raise ValueError("Generated CV section order does not match the requested order")
